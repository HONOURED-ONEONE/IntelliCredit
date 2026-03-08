from pathlib import Path
from loguru import logger
import re
import xml.sax.saxutils as saxutils


def _strip_markdown(text: str) -> str:
    """
    Remove a minimal subset of Markdown safely for PDF generation.
    This avoids malformed inline HTML that ReportLab can choke on.
    """
    if not text:
        return ""

    # Remove heading markers
    text = re.sub(r"^#{1,6}\s*", "", text)

    # Convert **bold** -> plain text
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)

    # Convert *italic* -> plain text
    text = re.sub(r"\*(.+?)\*", r"\1", text)

    # Remove inline code ticks
    text = re.sub(r"`(.+?)`", r"\1", text)

    return text.strip()


def _safe_pdf_text(text: str) -> str:
    """
    Prepare plain text for ReportLab Paragraph safely.
    We intentionally avoid injecting HTML tags like <b>...</b>
    because malformed tag sequences caused the prior PDF parse errors.
    """
    return saxutils.escape(_strip_markdown(text))


def cam_to_docx(job_dir: Path) -> Path:
    cam_md_path = job_dir / "decision_engine" / "cam.md"
    out_path = job_dir / "decision_engine" / "cam.docx"

    if not cam_md_path.exists():
        return None

    try:
        from docx import Document

        doc = Document()
        doc.add_heading("Credit Approval Memo (CAM)", 0)

        with open(cam_md_path, "r", encoding="utf-8") as f:
            content = f.read()

        for line in content.split("\n"):
            line = line.strip()
            if not line:
                continue

            if line.startswith("# "):
                doc.add_heading(_strip_markdown(line[2:]), level=1)
            elif line.startswith("## "):
                doc.add_heading(_strip_markdown(line[3:]), level=2)
            elif line.startswith("- "):
                doc.add_paragraph(_strip_markdown(line[2:]), style="List Bullet")
            else:
                doc.add_paragraph(_strip_markdown(line))

        doc.save(str(out_path))
        return out_path

    except Exception as e:
        logger.error(f"Failed to create DOCX: {e}")
        return None


def cam_to_pdf(job_dir: Path) -> Path:
    cam_md_path = job_dir / "decision_engine" / "cam.md"
    out_path = job_dir / "decision_engine" / "cam.pdf"

    if not cam_md_path.exists():
        return None

    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet

        doc = SimpleDocTemplate(str(out_path), pagesize=letter)
        styles = getSampleStyleSheet()
        story = []

        with open(cam_md_path, "r", encoding="utf-8") as f:
            content = f.read()

        for line in content.split("\n"):
            line = line.strip()

            if not line:
                story.append(Spacer(1, 12))
                continue

            if line.startswith("# "):
                text = _safe_pdf_text(line[2:])
                story.append(Paragraph(text, styles["Heading1"]))
            elif line.startswith("## "):
                text = _safe_pdf_text(line[3:])
                story.append(Paragraph(text, styles["Heading2"]))
            elif line.startswith("- "):
                text = _safe_pdf_text(line[2:])
                story.append(Paragraph(text, styles["Normal"], bulletText="•"))
            else:
                text = _safe_pdf_text(line)
                story.append(Paragraph(text, styles["Normal"]))

        doc.build(story)
        return out_path

    except Exception as e:
        logger.error(f"Failed to create PDF: {e}")
        return None